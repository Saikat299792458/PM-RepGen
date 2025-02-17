from docx import Document
import pandas as pd
from extractors import extractors

def main():
    doc = Document("test.docx") # Load the .docx file
    data = [] # Initialize a list to store the extracted data
    ext = extractors() # Create extractor instance
    ext.refresh() # Create ext.data
    prev_tab = 0
    
    for i, table in enumerate(doc.tables):
        print(f"\rProcessing Tables...({i+1}/{len(doc.tables)})", end="")
        cell1 = table.rows[0].cells[0].text.strip()
        if cell1.startswith("Ref") or cell1.startswith("m~Î:"):
            cur_tab = 1
        elif cell1.startswith("mvavib"):
            cur_tab = 2
        elif cell1.startswith("1- evwn"):
            cur_tab = 3
        elif cell1.startswith("4-D`i"):
            cur_tab = 4
        elif cell1.startswith("we:`"):
            cur_tab = 5
        
        if cur_tab < prev_tab:
            data.append(ext.data)
            ext.refresh()
        try:
            exec(f"ext.table{cur_tab}(table, {i+1})")
        except Exception as e:
            print(e)
        prev_tab = cur_tab
    data.append(ext.data)
            

    #for i in data:
     #   print(i)
    #### Create Excel file and Format
    columns = pd.MultiIndex.from_tuples([
    ("PM No","","",""),
    ("PM Date","","",""),
    ("Case No","","",""), 
    ("Ref","","",""), 
    ("Date","","",""),
    ("Name","","",""), 
    ("Sex","","",""), 
    ("Age","","",""), 
    ("Condition of the body", "Built","",""), 
    ("Condition of the body", "Rigor Mortis","All Over",""), 
    ("Condition of the body", "Rigor Mortis","Partial",""), 
    ("Condition of the body", "Decomposition","",""),
    ("Tattoo Marks", "", "", ""),
    ("Bladder", "", "", ""),
    ("Genitalia", "Male", "", ""),
    ("Genitalia", "Female", "Uterus", ""),
    ("Genitalia", "Female", "Vagina", ""),
    ("Genitalia", "Female", "Hymen", ""),
    ("Viscera Sent", "Yes", "Chemical Analysis", "Routine"),
    ("Viscera Sent", "Yes", "Chemical Analysis", "Others"),
    ("Viscera Sent", "Yes", "Histopathological Examination", ""),
    ("Viscera Sent", "Yes", "Microscopic Examination", "HVS"),
    ("Viscera Sent", "Yes", "Microscopic Examination", "Others"),
    ("Viscera Sent", "Yes", "DNA Analysis", ""),
    ("Viscera Sent", "No", "", ""),
    ("Mode of Death", "", "", ""),
    ("Manner of Death","Natural","",""),
    ("Manner of Death","Unnatural","Suicide",""),
    ("Manner of Death","Unnatural","Homicide",""),
    ("Manner of Death","Unnatural","Accident",""),
    ("Cause of Death", "Hanging", "", ""),
    ("Cause of Death", "Strangulation", "", ""),
    ("Cause of Death", "Suffocation", "", ""),
    ("Cause of Death", "Drowning", "", ""),
    ("Cause of Death", "Poison", "", ""),
    ("Cause of Death", "Burn", "", ""),
    ("Cause of Death", "Electrocution", "", ""),
    ("Cause of Death", "Injury", "Blunt Weapon", "Abrasion"),
    ("Cause of Death", "Injury", "Blunt Weapon", "Bruise"),
    ("Cause of Death", "Injury", "Blunt Weapon", "Laceration"),
    ("Cause of Death", "Injury", "Blunt Weapon", "Fracture"),
    ("Cause of Death", "Injury", "Blunt Weapon", "Dislocation"),
    ("Cause of Death", "Injury", "Sharp Weapon", "Incised Wound"),
    ("Cause of Death", "Injury", "Sharp Weapon", "Chop Wound"),
    ("Cause of Death", "Injury", "Sharp Weapon", "Stab Wound"),
    ("Cause of Death", "Injury", "Self Inflicted Wound", ""),
    ("Cause of Death", "Injury", "Fabricated Wound", ""),
    ("Cause of Death", "Injury", "Defence Wound", ""),
    ("Cause of Death", "Injury", "Surgical Wound", ""),
    ("Cause of Death", "Injury", "Gun Shot Injury", ""),
    ("Cause of Death", "Fall from Height", "", ""),
    ("Cause of Death", "Road Traffic Accident", "", "")
])

    data = [list(d.values()) for d in data]
    #print(data)
    filename = "PM_Basic_Info.xlsx"
    df = pd.DataFrame(data, columns=columns)
    df.index = range(1, len(df) + 1)
    df.index.name="Sl No."
    #df.insert(0, "Sl No.", range(1, len(df) + 1)) # Add Sl No. as the primary key
    #df.to_excel(filename) # Save the DataFrame to an Excel file
    writer = pd.ExcelWriter(filename, engine="xlsxwriter")
    df.to_excel(writer, sheet_name="Sheet1")

    # Get workbook and worksheet objects
    workbook = writer.book
    worksheet = writer.sheets["Sheet1"]

    worksheet.autofit()
    writer.close()

if __name__ == "__main__":
    main()