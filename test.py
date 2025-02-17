from docx import Document
from bijoy2unicode import converter

# Load the .docx file
doc = Document("test.docx")
test = converter.Unicode()

table = doc.tables[2]
row1 = table.columns[2]
cell = row1.cells[1]
toprint = test.convertBijoyToUnicode(cell.text.strip())
print(toprint)
print(cell.text.strip())


# Iterate through the tables in the document
    #for table in doc.tables:
    #    prev_tab = prev_tab + 1 if prev_tab < 5 else 1
    #    try:
    #        exec(f"ext.table{prev_tab}(table)")
    #    except Exception as e:
    #        print(e)
    #    if prev_tab == 5:
    #        data.append(ext.data)     
    #        ext.refresh()